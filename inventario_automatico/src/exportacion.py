from fpdf import FPDF
from odf.opendocument import OpenDocumentText, OpenDocumentSpreadsheet
from odf.text import P
from odf.table import Table, TableRow, TableCell
from docx import Document

def exportar_a_excel(df, ruta_salida):
    """ Guarda el DataFrame en un archivo Excel (.xlsx) """
    try:
        if df is None or df.empty:
            raise ValueError("No se puede exportar a Excel, el DataFrame está vacío.")
        df.to_excel(ruta_salida, index=False)
        print(f"✅ Reporte exportado a Excel: {ruta_salida}")
    except Exception as e:
        print(f"⚠️ Error al exportar a Excel: {e}")


def exportar_a_ods(df, ruta_salida):
    """ Guarda el DataFrame en un archivo de hoja de cálculo OpenDocument (.ods) """
    try:
        if df is None or df.empty:
            raise ValueError("No se puede exportar a ODS, el DataFrame está vacío.")

        doc = OpenDocumentSpreadsheet()
        tabla = Table()

        # Agregar encabezados
        encabezados = TableRow()
        for col in df.columns:
            celda = TableCell()
            celda.addElement(P(text=col))
            encabezados.addElement(celda)
        tabla.addElement(encabezados)

        # Agregar filas de datos
        for _, row in df.iterrows():
            fila = TableRow()
            for cell in row:
                celda = TableCell()
                celda.addElement(P(text=str(cell)))
                fila.addElement(celda)
            tabla.addElement(fila)

        doc.spreadsheet.addElement(tabla)
        doc.save(ruta_salida)
        print(f"✅ Reporte exportado a ODS: {ruta_salida}")
    except Exception as e:
        print(f"⚠️ Error al exportar a ODS: {e}")


def exportar_a_odt(df, ruta_salida):
    """ Guarda el DataFrame en un archivo de texto OpenDocument (.odt) """
    try:
        if df is None or df.empty:
            raise ValueError("No se puede exportar a ODT, el DataFrame está vacío.")

        doc = OpenDocumentText()
        for col in df.columns:
            doc.text.addElement(P(text=col))

        for _, row in df.iterrows():
            doc.text.addElement(P(text=" | ".join(str(cell) for cell in row)))

        doc.save(ruta_salida)
        print(f"✅ Reporte exportado a ODT: {ruta_salida}")
    except Exception as e:
        print(f"⚠️ Error al exportar a ODT: {e}")


def exportar_a_pdf(df, ruta_salida):
    """ Exporta el DataFrame a PDF """
    try:
        if df is None or df.empty:
            raise ValueError("No se puede exportar a PDF, el DataFrame está vacío.")

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=10)

        pdf.cell(200, 10, "Reporte de Inventario", ln=True, align='C')

        for col in df.columns:
            pdf.cell(40, 10, col, border=1)
        pdf.ln()

        for _, row in df.iterrows():
            for cell in row:
                pdf.cell(40, 10, str(cell), border=1)
            pdf.ln()

        pdf.output(ruta_salida)
        print(f"✅ Reporte exportado a PDF: {ruta_salida}")
    except Exception as e:
        print(f"⚠️ Error al exportar a PDF: {e}")


def exportar_a_docx(df, ruta_salida):
    """ Exporta el DataFrame a un documento de Word (.docx) """
    try:
        if df is None or df.empty:
            raise ValueError("No se puede exportar a DOCX, el DataFrame está vacío.")

        doc = Document()
        doc.add_heading("Reporte de Inventario", level=1)

        tabla = doc.add_table(rows=1, cols=len(df.columns))
        tabla.style = "Table Grid"

        # Agregar encabezados
        hdr_cells = tabla.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = col

        # Agregar filas
        for _, row in df.iterrows():
            fila = tabla.add_row().cells
            for i, cell in enumerate(row):
                fila[i].text = str(cell)

        doc.save(ruta_salida)
        print(f"✅ Reporte exportado a DOCX: {ruta_salida}")
    except Exception as e:
        print(f"⚠️ Error al exportar a DOCX: {e}")


def exportar_a_txt(df, ruta_salida):
    """ Guarda el DataFrame en un archivo de texto plano """
    try:
        if df is None or df.empty:
            raise ValueError("No se puede exportar a TXT, el DataFrame está vacío.")
        df.to_csv(ruta_salida, sep="\t", index=False)
        print(f"✅ Reporte exportado a TXT: {ruta_salida}")
    except Exception as e:
        print(f"⚠️ Error al exportar a TXT: {e}")
